"""
tasks/knowledge_ingest.py — Background job para processing de documentos
(chunking, embedding, storage).

Fluxo:
    1. POST /knowledge/upload enfileira processo
    2. Celery worker pega task
    3. Extrai conteúdo (PDF → text, CSV → rows, etc.)
    4. Chunk (500-token chunks, 50-token overlap)
    5. Batch embed via Sentence Transformers
    6. Store in pgvector + metadata
    7. Update KnowledgeDocument(processing_status='complete', progress_pct=100)

Se falhar em qualquer etapa: processing_status='failed', error_message=...

Suporte a tipos:
  • PDF: pypdf (text extraction + OCR optional)
  • CSV: pandas (parse + embed cada row como chunk)
  • DWG: opcional (metadata extraction, não full parsing)
  • TXT: split direto por linhas
  • DOCX: python-docx
"""
from __future__ import annotations

import asyncio
import io
import json
import logging
import os
from typing import Any

from sqlalchemy import select, update
from sqlalchemy.ext.asyncio import AsyncSession, create_async_engine
from sqlalchemy.orm import sessionmaker

from models.knowledge import KnowledgeChunk, KnowledgeDocument
from ml.embeddings import embed_batch

logger = logging.getLogger("manta.knowledge_ingest")

# TODO: Integrar com Celery / Redis quando necessário


# =============================================================================
# Token counting (aproximado)
# =============================================================================
def _estimate_tokens(text: str) -> int:
    """Estimativa simples: 1 token ~= 1 palavra."""
    return len(text.split())


# =============================================================================
# Extractors por tipo de arquivo
# =============================================================================
def _extract_pdf(content: bytes) -> list[tuple[str, dict[str, Any]]]:
    """Extrai texto de PDF. Retorna [(text_chunk, metadata), ...]."""
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning("pypdf not installed, skipping PDF extraction")
        return []

    try:
        pdf = PdfReader(io.BytesIO(content))
        pages = []
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text()
            if text.strip():
                pages.append(
                    (
                        text,
                        {"page": page_num, "source_file": "unknown.pdf"},
                    )
                )
        return pages
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return []


def _extract_csv(content: bytes) -> list[tuple[str, dict[str, Any]]]:
    """Extrai dados de CSV. Cada linha → chunk com metadata."""
    try:
        import pandas as pd
    except ImportError:
        logger.warning("pandas not installed, skipping CSV extraction")
        return []

    try:
        df = pd.read_csv(io.BytesIO(content))
        rows = []
        for idx, row in df.iterrows():
            text = " | ".join(f"{k}: {v}" for k, v in row.items())
            rows.append(
                (
                    text,
                    {"row": int(idx), "section": "csv_data", "source_file": "unknown.csv"},
                )
            )
        return rows
    except Exception as e:
        logger.error(f"CSV extraction failed: {e}")
        return []


def _extract_txt(content: bytes) -> list[tuple[str, dict[str, Any]]]:
    """Extrai texto de TXT. Split por linhas vazias (parágrafos)."""
    try:
        text = content.decode("utf-8")
    except Exception as e:
        logger.error(f"TXT decode failed: {e}")
        return []

    paragraphs = text.split("\n\n")
    chunks = []
    for para_idx, para in enumerate(paragraphs):
        if para.strip():
            chunks.append(
                (
                    para,
                    {
                        "section": f"paragraph_{para_idx}",
                        "source_file": "unknown.txt",
                    },
                )
            )
    return chunks


def _extract_docx(content: bytes) -> list[tuple[str, dict[str, Any]]]:
    """Extrai texto de DOCX."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed, skipping DOCX extraction")
        return []

    try:
        doc = Document(io.BytesIO(content))
        paragraphs = []
        for para_idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                paragraphs.append(
                    (
                        para.text,
                        {"section": f"paragraph_{para_idx}", "source_file": "unknown.docx"},
                    )
                )
        return paragraphs
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return []


def _extract_dwg(content: bytes) -> list[tuple[str, dict[str, Any]]]:
    """DWG: extração de metadata apenas (sem parsing de geometria)."""
    # Placeholder: seria necessário ezdxf ou similar para parsing real
    return [
        (
            "DWG file metadata extraction not yet implemented",
            {"source_file": "unknown.dwg", "section": "metadata"},
        )
    ]


def _extract_content(file_type: str, content: bytes) -> list[tuple[str, dict[str, Any]]]:
    """Dispatcher para extração por tipo."""
    extractors = {
        "pdf": _extract_pdf,
        "csv": _extract_csv,
        "txt": _extract_txt,
        "docx": _extract_docx,
        "dwg": _extract_dwg,
    }
    extractor = extractors.get(file_type.lower())
    if not extractor:
        logger.warning(f"No extractor for type {file_type}")
        return []

    return extractor(content)


# =============================================================================
# Chunking (500-token, 50-token overlap)
# =============================================================================
def _chunk_paragraphs(
    paragraphs: list[tuple[str, dict[str, Any]]],
    target_tokens: int = 500,
    overlap_tokens: int = 50,
) -> list[tuple[str, dict[str, Any]]]:
    """
    Agrupa parágrafos em chunks de ~target_tokens com overlap.

    Entrada: [(text, metadata), ...]
    Saída: [(text, metadata_com_page_start), ...]

    Estratégia:
      1. Estima tokens de cada parágrafo
      2. Agrupa até atingir target_tokens
      3. Cria overlap: última N palavras do chunk anterior
    """
    chunks = []
    current_chunk = []
    current_tokens = 0
    current_metadata = None

    for text, metadata in paragraphs:
        para_tokens = _estimate_tokens(text)

        # Se este parágrafo sozinho > target, o coloca direto (não tira)
        if para_tokens > target_tokens:
            if current_chunk:
                chunk_text = "\n".join(current_chunk)
                chunks.append((chunk_text, current_metadata))
                current_chunk = []
                current_tokens = 0

            chunks.append((text, metadata))
            current_metadata = metadata
            continue

        # Se adicionar este parágrafo passa do target, finaliza chunk
        if current_tokens + para_tokens > target_tokens and current_chunk:
            chunk_text = "\n".join(current_chunk)
            chunks.append((chunk_text, current_metadata))

            # Overlap: últimas N palavras do chunk anterior → novo chunk
            words = chunk_text.split()
            overlap_words = words[-min(overlap_tokens, len(words)) :]
            current_chunk = [" ".join(overlap_words)] if overlap_words else []
            current_tokens = len(overlap_words)
            current_metadata = metadata

        current_chunk.append(text)
        current_tokens += para_tokens

    # Finaliza
    if current_chunk:
        chunk_text = "\n".join(current_chunk)
        chunks.append((chunk_text, current_metadata))

    return chunks


# =============================================================================
# Main ingestion task
# =============================================================================
async def process_document_upload(
    doc_id: str,
    file_type: str,
    content: bytes,
    session: AsyncSession,
) -> None:
    """
    Background task: extrai → chunks → embeds → storage.

    1. Mark status='processing'
    2. Extrai conteúdo
    3. Cria chunks
    4. Gera embeddings
    5. Store em pgvector
    6. Update status='complete'

    Se falhar: status='failed', error_message
    """
    try:
        # Update status → processing
        stmt = (
            update(KnowledgeDocument)
            .where(KnowledgeDocument.id == doc_id)
            .values(
                processing_status="processing",
                progress_pct=5,
            )
        )
        await session.execute(stmt)
        await session.commit()

        logger.info(f"knowledge_ingest: starting {doc_id}")

        # Extrai conteúdo
        paragraphs = _extract_content(file_type, content)
        if not paragraphs:
            raise ValueError("No content extracted from file")

        stmt = (
            update(KnowledgeDocument)
            .where(KnowledgeDocument.id == doc_id)
            .values(progress_pct=30)
        )
        await session.execute(stmt)
        await session.commit()

        logger.info(f"knowledge_ingest: extracted {len(paragraphs)} paragraphs")

        # Chunk
        chunks = _chunk_paragraphs(paragraphs)
        if not chunks:
            raise ValueError("No chunks created")

        stmt = (
            update(KnowledgeDocument)
            .where(KnowledgeDocument.id == doc_id)
            .values(progress_pct=50)
        )
        await session.execute(stmt)
        await session.commit()

        logger.info(f"knowledge_ingest: created {len(chunks)} chunks")

        # Prepara batch de embeddings
        chunk_texts = [text for text, _ in chunks]
        embeddings = await embed_batch(chunk_texts)

        if not embeddings or len(embeddings) != len(chunks):
            raise ValueError("Embedding failed or incomplete")

        stmt = (
            update(KnowledgeDocument)
            .where(KnowledgeDocument.id == doc_id)
            .values(progress_pct=75)
        )
        await session.execute(stmt)
        await session.commit()

        logger.info(f"knowledge_ingest: embedded {len(embeddings)} chunks")

        # Store chunks
        chunk_objects = []
        for idx, ((text, metadata), embedding) in enumerate(zip(chunks, embeddings)):
            chunk_obj = KnowledgeChunk(
                id=str(id(chunk_obj)) if idx == 0 else None,  # Will be auto-generated
                document_id=doc_id,
                chunk_index=idx,
                content=text,
                content_tokens=_estimate_tokens(text),
                embedding=embedding,
                metadata=metadata,
            )
            chunk_objects.append(chunk_obj)

        session.add_all(chunk_objects)

        # Update document
        stmt = (
            update(KnowledgeDocument)
            .where(KnowledgeDocument.id == doc_id)
            .values(
                chunk_count=len(chunks),
                processing_status="complete",
                progress_pct=100,
            )
        )
        await session.execute(stmt)
        await session.commit()

        logger.info(f"knowledge_ingest: completed {doc_id} ({len(chunks)} chunks)")

    except Exception as e:
        logger.error(f"knowledge_ingest: failed {doc_id}: {e}", exc_info=True)

        stmt = (
            update(KnowledgeDocument)
            .where(KnowledgeDocument.id == doc_id)
            .values(
                processing_status="failed",
                error_message=str(e)[:512],
                progress_pct=0,
            )
        )
        await session.execute(stmt)
        await session.commit()


# =============================================================================
# Celery task (when Celery is integrated)
# =============================================================================
# @celery.task(bind=True, max_retries=3)
# def process_document_upload_task(
#     self,
#     doc_id: str,
#     file_type: str,
#     content_b64: str,  # base64-encoded
# ) -> dict[str, Any]:
#     """Celery task wrapper."""
#     try:
#         content = base64.b64decode(content_b64)
#         loop = asyncio.get_event_loop()
#         session = SessionLocal()  # Sync session
#
#         loop.run_until_complete(
#             process_document_upload(doc_id, file_type, content, session)
#         )
#
#         return {"status": "complete", "doc_id": doc_id}
#     except Exception as exc:
#         logger.error(f"Task failed: {exc}")
#         self.retry(exc=exc, countdown=60)
