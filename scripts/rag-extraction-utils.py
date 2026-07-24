#!/usr/bin/env python3
"""
RAG Extraction Utilities
Ferramentas para extração, chunkarização e validação de documentos
"""

import os
import re
import json
import hashlib
from pathlib import Path
from typing import List, Dict, Tuple
from datetime import datetime

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import openpyxl
except ImportError:
    openpyxl = None


class DocumentExtractor:
    """Extrator de conteúdo de documentos (PDF, DOCX, XLSX, TXT)"""

    def __init__(self, encoding='utf-8'):
        self.encoding = encoding

    def extract_pdf(self, filepath: str) -> str:
        """Extrair texto de PDF"""
        if not PyPDF2:
            raise ImportError("PyPDF2 não instalado. Execute: pip install PyPDF2")

        text = []
        try:
            with open(filepath, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text.append(page.extract_text())
        except Exception as e:
            raise ValueError(f"Erro ao extrair PDF {filepath}: {e}")

        return '\n'.join(text)

    def extract_docx(self, filepath: str) -> str:
        """Extrair texto de DOCX"""
        if not DocxDocument:
            raise ImportError("python-docx não instalado. Execute: pip install python-docx")

        try:
            doc = DocxDocument(filepath)
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            return '\n'.join(text)
        except Exception as e:
            raise ValueError(f"Erro ao extrair DOCX {filepath}: {e}")

    def extract_xlsx(self, filepath: str) -> str:
        """Extrair texto de XLSX"""
        if not openpyxl:
            raise ImportError("openpyxl não instalado. Execute: pip install openpyxl")

        try:
            wb = openpyxl.load_workbook(filepath)
            text = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                text.append(f"\n=== Sheet: {sheet} ===\n")
                for row in ws.iter_rows(values_only=True):
                    text.append('\t'.join(str(cell) if cell else '' for cell in row))
            return '\n'.join(text)
        except Exception as e:
            raise ValueError(f"Erro ao extrair XLSX {filepath}: {e}")

    def extract_txt(self, filepath: str) -> str:
        """Extrair texto de TXT"""
        try:
            with open(filepath, 'r', encoding=self.encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            # Tentar com encoding diferente
            try:
                with open(filepath, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                raise ValueError(f"Erro ao extrair TXT {filepath}: {e}")

    def extract(self, filepath: str) -> str:
        """Extrair conteúdo baseado no tipo de arquivo"""
        filepath = Path(filepath)

        if not filepath.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {filepath}")

        suffix = filepath.suffix.lower()

        if suffix == '.pdf':
            return self.extract_pdf(str(filepath))
        elif suffix == '.docx':
            return self.extract_docx(str(filepath))
        elif suffix == '.xlsx':
            return self.extract_xlsx(str(filepath))
        elif suffix == '.txt':
            return self.extract_txt(str(filepath))
        else:
            raise ValueError(f"Tipo de arquivo não suportado: {suffix}")


class ContentCleaner:
    """Limpeza e normalização de conteúdo extraído"""

    @staticmethod
    def clean(text: str) -> str:
        """Limpar e normalizar texto"""
        # Remover múltiplos espaços em branco
        text = re.sub(r'\s+', ' ', text)

        # Remover linhas vazias múltiplas
        text = re.sub(r'\n\n+', '\n', text)

        # Remover caracteres de controle (exceto \n)
        text = ''.join(char for char in text if char == '\n' or not (ord(char) < 32 and char != '\n'))

        # Trim
        text = text.strip()

        return text


class DocumentChunker:
    """Dividir documentos em chunks com metadados"""

    def __init__(self, chunk_size: int = 1000, overlap: int = 100):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def create_chunks(self, text: str, metadata: Dict = None) -> List[Dict]:
        """Criar chunks de um documento"""
        if metadata is None:
            metadata = {}

        chunks = []
        chunk_index = 0

        # Dividir por parágrafos primeiro
        paragraphs = text.split('\n')
        current_chunk = []
        current_length = 0

        for para in paragraphs:
            para_length = len(para) + 1  # +1 para newline

            if current_length + para_length > self.chunk_size and current_chunk:
                # Salvar chunk atual
                chunk_text = '\n'.join(current_chunk).strip()
                if chunk_text:
                    chunks.append({
                        'chunk_index': chunk_index,
                        'content': chunk_text,
                        'length': len(chunk_text),
                        **metadata
                    })
                    chunk_index += 1

                # Iniciar novo chunk com overlap
                if len(current_chunk) > 1:
                    # Manter últimas linhas para overlap
                    overlap_paras = []
                    overlap_length = 0
                    for p in reversed(current_chunk):
                        if overlap_length + len(p) <= self.overlap:
                            overlap_paras.insert(0, p)
                            overlap_length += len(p) + 1
                        else:
                            break
                    current_chunk = overlap_paras
                    current_length = overlap_length
                else:
                    current_chunk = []
                    current_length = 0

            current_chunk.append(para)
            current_length += para_length

        # Salvar último chunk
        if current_chunk:
            chunk_text = '\n'.join(current_chunk).strip()
            if chunk_text:
                chunks.append({
                    'chunk_index': chunk_index,
                    'content': chunk_text,
                    'length': len(chunk_text),
                    **metadata
                })

        return chunks


class MetadataExtractor:
    """Extrair metadados de documentos"""

    @staticmethod
    def extract_metadata(filepath: str, collection: str, segment: str) -> Dict:
        """Extrair metadados do documento"""
        file_path = Path(filepath)

        metadata = {
            'document_id': f"doc-{datetime.now().strftime('%Y-%m-%d')}-{file_path.stem}",
            'filename': file_path.name,
            'collection_prefix': collection,
            'segment': segment,
            'source_url': str(filepath),
            'source_type': 'file',
            'file_size_kb': file_path.stat().st_size / 1024,
            'file_format': file_path.suffix.lower(),
            'extracted_at': datetime.now().isoformat(),
        }

        return metadata


class RAGValidator:
    """Validação de chunks com aluci-guard"""

    def __init__(self, confidence_threshold: float = 0.85):
        self.confidence_threshold = confidence_threshold

    def validate(self, chunk_text: str, context: Dict = None) -> Dict:
        """
        Validar chunk com aluci-guard

        Em produção, isso chamaria a API do aluci-guard.
        Por enquanto, simular validação baseada em heurísticas.
        """
        if context is None:
            context = {}

        confidence_score = self._calculate_confidence(chunk_text, context)

        validation_result = {
            'valid': confidence_score >= self.confidence_threshold,
            'confidence_score': round(confidence_score, 2),
            'validation_status': 'validated' if confidence_score >= self.confidence_threshold else 'pending',
            'issues': self._detect_issues(chunk_text),
            'validated_at': datetime.now().isoformat()
        }

        return validation_result

    @staticmethod
    def _calculate_confidence(text: str, context: Dict) -> float:
        """Calcular confidence score (heurístico)"""
        score = 0.95  # Baseline alto

        # Penalidade por características suspeitas
        if len(text) < 100:
            score -= 0.1  # Chunk muito pequeno

        if re.search(r'(TODO|FIXME|XXX|HACK)', text, re.IGNORECASE):
            score -= 0.15  # Contém marcadores de rascunho

        if re.search(r'(fictional|example|hypothetical)', text, re.IGNORECASE):
            score -= 0.2  # Contém indicadores de ficção

        # Bônus por características boas
        if re.search(r'(Lei|NBR|ABNT|ANEEL|SNIS)', text):
            score += 0.05  # Referências a normas/órgãos

        # Garantir que score está entre 0 e 1
        return max(0.0, min(1.0, score))

    @staticmethod
    def _detect_issues(text: str) -> List[str]:
        """Detectar possíveis problemas no texto"""
        issues = []

        # Verificar URLs suspeitas
        urls = re.findall(r'https?://[^\s]+', text)
        for url in urls:
            if 'example.com' in url or 'fake' in url.lower():
                issues.append(f"URL suspeita: {url}")

        # Verificar normas fabricadas
        norms = re.findall(r'(NBR|ABNT)\s+(\d+[.-]\d+)', text)
        for norm_type, norm_number in norms:
            if not RAGValidator._is_valid_norm(norm_type, norm_number):
                issues.append(f"Norma potencialmente fabricada: {norm_type} {norm_number}")

        return issues

    @staticmethod
    def _is_valid_norm(norm_type: str, norm_number: str) -> bool:
        """Verificar se norma é válida (muito simples)"""
        # Em produção, isso consultaria um banco de dados de normas
        return True  # Aceitar por enquanto


def process_document(filepath: str, collection: str, segment: str,
                     chunk_size: int = 1000) -> Tuple[List[Dict], Dict]:
    """
    Pipeline completo: extrair → limpar → chunkarizar → validar

    Returns:
        (chunks, metadata)
    """
    # 1. Extrair
    extractor = DocumentExtractor()
    raw_text = extractor.extract(filepath)

    # 2. Limpar
    cleaner = ContentCleaner()
    cleaned_text = cleaner.clean(raw_text)

    # 3. Extrair metadados
    metadata_extractor = MetadataExtractor()
    metadata = metadata_extractor.extract_metadata(filepath, collection, segment)

    # 4. Chunkarizar
    chunker = DocumentChunker(chunk_size=chunk_size)
    chunks = chunker.create_chunks(cleaned_text, metadata={'collection_prefix': collection, 'segment': segment})

    # 5. Validar cada chunk
    validator = RAGValidator(confidence_threshold=0.85)
    for chunk in chunks:
        validation = validator.validate(chunk['content'], context={'collection': collection, 'segment': segment})
        chunk.update(validation)

    return chunks, metadata


if __name__ == '__main__':
    import sys

    if len(sys.argv) < 3:
        print("Uso: python rag-extraction-utils.py <filepath> <collection> <segment>")
        print("Exemplo: python rag-extraction-utils.py doc.pdf san: S8")
        sys.exit(1)

    filepath = sys.argv[1]
    collection = sys.argv[2]
    segment = sys.argv[3]

    try:
        chunks, metadata = process_document(filepath, collection, segment)

        print(json.dumps({
            'metadata': metadata,
            'chunk_count': len(chunks),
            'chunks': chunks[:3]  # Mostrar primeiros 3 chunks
        }, indent=2, ensure_ascii=False))
    except Exception as e:
        print(f"Erro: {e}", file=sys.stderr)
        sys.exit(1)
